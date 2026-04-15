#!/usr/bin/env python
"""
tools/md_to_docx.py — Convertitore markdown → docx per la relazione d'esame.

Parser supportato (sottoinsieme ragionevole di GFM):
  - Headings: # H1, ## H2, ### H3, #### H4
  - Paragrafi
  - Liste puntate (- / *)
  - Liste numerate (1.)
  - Tabelle GFM (| col | col |)
  - Code blocks (```)
  - Regole orizzontali (---) → page break
  - Grassetto **...**, corsivo *...*, code `...`

Uso:
  python tools/md_to_docx.py [input.md] [output.docx]
Default: input = docs/relazione.md, output = docs/relazione.docx
"""
import re
import sys
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL

INLINE_BOLD = re.compile(r"\*\*([^*]+)\*\*")
INLINE_ITALIC = re.compile(r"(?<!\*)\*([^*]+)\*(?!\*)")
INLINE_CODE = re.compile(r"`([^`]+)`")


def add_inline_runs(paragraph, text):
    """Aggiunge run al paragrafo interpretando **bold**, *italic*, `code`."""
    pos = 0
    pattern = re.compile(r"(\*\*[^*]+\*\*|`[^`]+`|(?<!\*)\*[^*]+\*(?!\*))")
    for m in pattern.finditer(text):
        if m.start() > pos:
            paragraph.add_run(text[pos:m.start()])
        chunk = m.group(0)
        if chunk.startswith("**") and chunk.endswith("**"):
            run = paragraph.add_run(chunk[2:-2])
            run.bold = True
        elif chunk.startswith("`") and chunk.endswith("`"):
            run = paragraph.add_run(chunk[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0x5B, 0x3F, 0xD9)
        elif chunk.startswith("*") and chunk.endswith("*"):
            run = paragraph.add_run(chunk[1:-1])
            run.italic = True
        pos = m.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def parse_table_row(line):
    """Parse riga tabella GFM → lista di celle."""
    cells = [c.strip() for c in line.strip().strip("|").split("|")]
    return cells


def is_table_separator(line):
    return bool(re.fullmatch(r"\|?[\s:\-|]+\|?", line.strip())) and "-" in line


def flush_table(doc, rows):
    if not rows:
        return
    header = rows[0]
    body = rows[2:] if len(rows) > 1 and is_table_separator("|".join(rows[1])) else rows[1:]
    # Semplificazione: se row 1 è separator, body parte da row 2
    if len(rows) > 1 and all(set(c) <= set("-: ") for c in rows[1]):
        header = rows[0]
        body = rows[2:]
    else:
        header = rows[0]
        body = rows[1:]

    cols = len(header)
    table = doc.add_table(rows=1 + len(body), cols=cols)
    table.style = "Light Grid Accent 1"

    # Header row
    for i, cell_text in enumerate(header):
        cell = table.rows[0].cells[i]
        p = cell.paragraphs[0]
        run = p.add_run(cell_text)
        run.bold = True

    for r_idx, row in enumerate(body, start=1):
        for c_idx in range(cols):
            text = row[c_idx] if c_idx < len(row) else ""
            cell = table.rows[r_idx].cells[c_idx]
            p = cell.paragraphs[0]
            add_inline_runs(p, text)


def flush_code_block(doc, code_lines):
    if not code_lines:
        return
    p = doc.add_paragraph()
    run = p.add_run("\n".join(code_lines))
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x1A, 0x16, 0x25)
    p.paragraph_format.left_indent = Cm(0.6)
    p.paragraph_format.space_after = Pt(8)


def convert(md_path, out_path):
    text = Path(md_path).read_text(encoding="utf-8")
    lines = text.split("\n")

    # Strip YAML front matter se presente
    if lines and lines[0].strip() == "---":
        for idx in range(1, len(lines)):
            if lines[idx].strip() == "---":
                lines = lines[idx + 1:]
                break

    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # Base style
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    in_code = False
    code_buf = []
    table_buf = []

    for line in lines:
        stripped = line.rstrip()

        # Code block
        if stripped.startswith("```"):
            if in_code:
                flush_code_block(doc, code_buf)
                code_buf = []
                in_code = False
            else:
                in_code = True
            continue

        if in_code:
            code_buf.append(stripped)
            continue

        # Flush table on non-table line
        if table_buf and not stripped.startswith("|"):
            flush_table(doc, table_buf)
            table_buf = []

        # Table row
        if stripped.startswith("|"):
            table_buf.append(parse_table_row(stripped))
            continue

        # Headings
        if stripped.startswith("# "):
            doc.add_heading(stripped[2:], level=1)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=3)
        elif stripped.startswith("#### "):
            doc.add_heading(stripped[5:], level=4)
        elif stripped == "---":
            doc.add_page_break()
        elif stripped.startswith("- ") or stripped.startswith("* "):
            p = doc.add_paragraph(style="List Bullet")
            add_inline_runs(p, stripped[2:])
        elif re.match(r"^\d+\.\s+", stripped):
            content = re.sub(r"^\d+\.\s+", "", stripped)
            p = doc.add_paragraph(style="List Number")
            add_inline_runs(p, content)
        elif stripped.startswith("> "):
            p = doc.add_paragraph(style="Intense Quote")
            add_inline_runs(p, stripped[2:])
        elif stripped == "":
            # blank line — skip (paragraphs naturally separated)
            pass
        else:
            p = doc.add_paragraph()
            add_inline_runs(p, stripped)

    # Flush remaining
    if table_buf:
        flush_table(doc, table_buf)

    Path(out_path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)
    return out_path


if __name__ == "__main__":
    md = sys.argv[1] if len(sys.argv) > 1 else "docs/relazione.md"
    docx = sys.argv[2] if len(sys.argv) > 2 else "docs/relazione.docx"
    out = convert(md, docx)
    size = Path(out).stat().st_size
    print(f"ok: {out} ({size:,} bytes)")
